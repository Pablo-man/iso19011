import os
import gradio as gr
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
import tempfile
import time
from datetime import datetime
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import io

# Cargar variables de entorno
load_dotenv()

# Configurar Google Generative AI (Gemini)
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)
modelo = genai.GenerativeModel("gemini-2.0-flash-exp")

# Ruta al documento Word de la ISO
ruta_docx = os.getenv("ISO_DOC_PATH", "ISO19011.docx")

def leer_docx(ruta):
    """Lee el contenido de un documento Word"""
    try:
        doc = Document(ruta)
        texto = "\n".join([p.text for p in doc.paragraphs])
        return texto
    except Exception as e:
        print(f"Error al leer el documento: {e}")
        return "No se pudo leer el documento ISO. Verifica la ruta y el formato."

# Leer el contenido del documento ISO
contenido_iso = leer_docx(ruta_docx)

# Historia del chat para mantener el contexto
historial_chat = []
caso_actual = None
planificacion_actual = None
solucion_actual = None

def generar_caso_estudio():
    """Genera un caso de estudio basado en el documento ISO"""
    prompt = f"""
    Basándote en la siguiente norma ISO:
    
    \"\"\"{contenido_iso}\"\"\"
    
    Genera un caso de estudio realista y detallado sobre una empresa ficticia 
    que necesita implementar esta norma ISO. El caso debe incluir:
    
    1. Descripción de la empresa (tamaño, sector, ubicación)
    2. Situación actual y problemas que enfrenta
    3. Necesidad específica de implementar la norma ISO
    4. Contexto adicional relevante para el caso
    
    El caso debe ser didáctico y plantear un escenario que permita aplicar
    los conceptos clave de la norma.
    """
    
    try:
        respuesta = modelo.generate_content(prompt).text
        global caso_actual
        caso_actual = respuesta
        return respuesta
    except Exception as e:
        return f"Error al generar el caso de estudio: {str(e)}"

def crear_planificacion(caso):
    """Crea una planificación para resolver el caso de estudio"""
    if not caso:
        return "Primero debes generar un caso de estudio."
    
    prompt = f"""
    Basándote en la siguiente norma ISO:
    
    \"\"\"{contenido_iso}\"\"\"
    
    Y considerando el siguiente caso de estudio:
    
    \"\"\"{caso}\"\"\"
    
    Desarrolla una planificación detallada para implementar la norma ISO en este caso.
    La planificación debe incluir:
    
    1. Objetivos claros y medibles
    2. Etapas de implementación con cronograma sugerido
    3. Recursos necesarios (humanos, técnicos, financieros)
    4. Indicadores para medir el progreso y éxito
    5. Consideraciones importantes y posibles obstáculos
    
    Estructura la planificación de manera clara y práctica para que pueda ser implementada.
    """
    
    try:
        respuesta = modelo.generate_content(prompt).text
        global planificacion_actual
        planificacion_actual = respuesta
        return respuesta
    except Exception as e:
        return f"Error al crear la planificación: {str(e)}"

def generar_solucion(caso):
    """Genera una solución ideal para el caso de estudio"""
    if not caso:
        return "Primero debes generar un caso de estudio."
    
    prompt = f"""
    Basándote en la siguiente norma ISO:
    
    \"\"\"{contenido_iso}\"\"\"
    
    Y considerando el siguiente caso de estudio:
    
    \"\"\"{caso}\"\"\"
    
    Desarrolla una solución completa y detallada que siga las mejores prácticas
    para implementar la norma ISO en este caso específico. Incluye:
    
    1. Estrategia de implementación paso a paso
    2. Documentación necesaria a desarrollar
    3. Procesos a establecer o modificar
    4. Mecanismos de seguimiento y evaluación
    5. Medidas para asegurar la mejora continua
    
    Esta solución servirá como referencia ideal para comparar con la respuesta
    del usuario.
    """
    
    try:
        respuesta = modelo.generate_content(prompt).text
        global solucion_actual
        solucion_actual = respuesta
        return respuesta
    except Exception as e:
        return f"Error al generar la solución: {str(e)}"

def comparar_respuestas(respuesta_usuario, caso):
    """Compara la respuesta del usuario con la solución generada por la IA"""
    if not caso or not solucion_actual:
        return "Error: No hay un caso de estudio o solución generada para comparar."
    
    prompt = f"""
    Basándote en la siguiente norma ISO:
    
    \"\"\"{contenido_iso}\"\"\"
    
    Compara la siguiente respuesta del usuario:
    
    \"\"\"{respuesta_usuario}\"\"\"
    
    Con esta solución de referencia:
    
    \"\"\"{solucion_actual}\"\"\"
    
    Para el siguiente caso de estudio:
    
    \"\"\"{caso}\"\"\"
    
    Realiza un análisis detallado que incluya:
    
    1. Porcentaje de similitud conceptual (expresado como un número del 0 al 100)
    2. Aspectos bien abordados en la respuesta del usuario
    3. Elementos faltantes o insuficientes
    4. Enfoques diferentes pero válidos (si los hay)
    5. Evalúa qué tanto se apega la respuesta del usuario al contenido del documento ISO
    
    Presenta los resultados de manera estructurada y constructiva.
    """
    
    try:
        resultado_comparacion = modelo.generate_content(prompt).text
        
        # Extraer el porcentaje de similitud para la gráfica
        try:
            lines = resultado_comparacion.split('\n')
            for line in lines:
                if '%' in line and ('similitud' in line.lower() or 'porcentaje' in line.lower()):
                    porcentaje_texto = ''.join(filter(lambda x: x.isdigit() or x == '.', line))
                    porcentaje = float(porcentaje_texto)
                    break
            else:
                # Si no se encuentra el porcentaje, asignar un valor predeterminado
                porcentaje = 50
        except:
            porcentaje = 50
        
        # Generar gráfica
        fig = generar_grafica_comparacion(porcentaje)
        
        return resultado_comparacion, fig
    except Exception as e:
        return f"Error al comparar las respuestas: {str(e)}", None

def generar_recomendaciones(respuesta_usuario, resultado_comparacion, caso):
    """Genera recomendaciones para la mejora continua"""
    if not caso:
        return "Error: No hay un caso de estudio para generar recomendaciones."
    
    prompt = f"""
    Basándote en la siguiente norma ISO:
    
    \"\"\"{contenido_iso}\"\"\"
    
    Y considerando:
    
    1. El caso de estudio: \"\"\"{caso}\"\"\"
    2. La respuesta del usuario: \"\"\"{respuesta_usuario}\"\"\"
    3. El resultado de la comparación: \"\"\"{resultado_comparacion}\"\"\"
    
    Proporciona recomendaciones detalladas y prácticas para mejorar la implementación
    propuesta por el usuario. Las recomendaciones deben:
    
    1. Ser específicas y accionables
    2. Centrarse en la mejora continua
    3. Referirse a secciones específicas de la norma ISO cuando sea relevante
    4. Ofrecer recursos adicionales o herramientas que podrían ser útiles
    5. Sugerir indicadores para medir el progreso de las mejoras
    
    Estructura las recomendaciones de manera clara y positiva, reconociendo los
    puntos fuertes del trabajo del usuario.
    """
    
    try:
        recomendaciones = modelo.generate_content(prompt).text
        return recomendaciones
    except Exception as e:
        return f"Error al generar recomendaciones: {str(e)}"

def generar_grafica_comparacion(porcentaje):
    """Genera una gráfica comparativa del porcentaje de similitud"""
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Datos para el gráfico
    categorias = ['Similitud con la solución ideal']
    valores = [porcentaje]
    colores = ['#3498db'] if porcentaje >= 70 else ['#f39c12'] if porcentaje >= 50 else ['#e74c3c']
    
    # Crear gráfico de barras
    bars = ax.bar(categorias, valores, color=colores, width=0.6)
    
    # Añadir etiquetas
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                f'{height:.1f}%', ha='center', va='bottom', fontsize=12, fontweight='bold')
    
    # Configurar ejes
    ax.set_ylim(0, 110)
    ax.set_ylabel('Porcentaje de similitud (%)', fontsize=12)
    ax.set_title('Comparación con la solución ideal', fontsize=14, fontweight='bold')
    
    # Añadir líneas de referencia
    ax.axhline(y=50, color='#f39c12', linestyle='--', alpha=0.5)
    ax.axhline(y=70, color='#3498db', linestyle='--', alpha=0.5)
    ax.text(1.01, 50, 'Aceptable (50%)', va='center', ha='left', fontsize=10, color='#f39c12')
    ax.text(1.01, 70, 'Bueno (70%)', va='center', ha='left', fontsize=10, color='#3498db')
    
    plt.tight_layout()
    return fig

def guardar_texto(texto, prefijo):
    """Guarda texto en un archivo temporal y devuelve la ruta"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{prefijo}_{timestamp}.txt"
    
    with open(filename, "w", encoding="utf-8") as f:
        f.write(texto)
    
    return filename

def generar_archivo_txt(texto, prefijo):
    """Genera un archivo TXT a partir del texto"""
    if not texto:
        return None
        
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{prefijo}_{timestamp}.txt"
    
    # Crear un objeto de archivo en memoria
    buffer = io.BytesIO()
    buffer.write(texto.encode('utf-8'))
    buffer.seek(0)
    
    return (filename, buffer)

def generar_archivo_docx(texto, prefijo):
    """Genera un archivo DOCX a partir del texto"""
    if not texto:
        return None
        
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{prefijo}_{timestamp}.docx"
    
    # Crear un documento Word
    doc = Document()
    doc.add_heading(f"{prefijo.capitalize()}", 0)
    
    # Añadir fecha y hora
    doc.add_paragraph(f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("\n")
    
    # Añadir el texto
    doc.add_paragraph(texto)
    
    # Guardar en memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return (filename, buffer)

def generar_pdf(texto, prefijo):
    """Placeholder para generación de PDF (requiere librería adicional)"""
    # Nota: Para implementar esto, necesitarías una librería como reportlab o fpdf
    # Esta es una implementación simplificada
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"La generación de PDF no está implementada. Usa TXT o DOCX."

def actualizar_chat(mensaje, historial=None):
    """Función para manejar el flujo del chat"""
    if historial is None:
        historial = []
    
    historial.append(mensaje)
    return historial

# Configuración de la interfaz Gradio
with gr.Blocks(theme=gr.themes.Soft(primary_hue="blue")) as demo:
    gr.Markdown("# Sistema de Casos de Estudio ISO 19011")
    gr.Markdown("Esta aplicación genera casos de estudio basados en normas ISO 10911, permite planificar su resolución, evaluar respuestas y proporcionar recomendaciones para la mejora continua.")
    
    with gr.Tab("Generador de Casos de Estudio"):
        with gr.Row():
            with gr.Column(scale=3):
                caso_btn = gr.Button("Generar Caso de Estudio", variant="primary")
                caso_output = gr.Textbox(
                    label="Caso de Estudio", 
                    placeholder="El caso de estudio aparecerá aquí...",
                    lines=15
                )
            
            with gr.Column(scale=1):
                with gr.Group():
                    gr.Markdown("### Descargar Caso")
                    download_caso_txt = gr.File(label="Descargar como TXT")
                    download_caso_docx = gr.File(label="Descargar como DOCX")
    
    with gr.Tab("Planificación"):
        with gr.Row():
            with gr.Column(scale=3):
                plan_btn = gr.Button("Crear Planificación", variant="primary")
                plan_output = gr.Textbox(
                    label="Planificación", 
                    placeholder="La planificación aparecerá aquí...",
                    lines=15
                )
            
            with gr.Column(scale=1):
                with gr.Group():
                    gr.Markdown("### Descargar Planificación")
                    download_plan_txt = gr.File(label="Descargar como TXT")
                    download_plan_docx = gr.File(label="Descargar como DOCX")
    
    with gr.Tab("Evaluación de Respuesta"):
        with gr.Row():
            respuesta_usuario = gr.Textbox(
                label="Tu Resolución al Caso", 
                placeholder="Escribe aquí tu propuesta de solución al caso de estudio...",
                lines=10
            )
            solucion_btn = gr.Button("Ver Solución de Referencia", variant="secondary")
        
        solucion_output = gr.Textbox(
            label="Solución de Referencia", 
            placeholder="La solución de referencia aparecerá aquí...",
            lines=10,
            visible=False
        )
        
        comparar_btn = gr.Button("Comparar Mi Respuesta", variant="primary")
        
        with gr.Row():
            with gr.Column(scale=2):
                resultado_comparacion = gr.Textbox(
                    label="Resultado de la Comparación", 
                    placeholder="El resultado de la comparación aparecerá aquí...",
                    lines=10
                )
            
            with gr.Column(scale=1):
                grafico_comparacion = gr.Plot(label="Gráfico de Comparación")
        
        with gr.Row():
            with gr.Column(scale=1):
                with gr.Group():
                    gr.Markdown("### Descargar Resolución")
                    download_sol_txt = gr.File(label="Descargar como TXT")
                    download_sol_docx = gr.File(label="Descargar como DOCX")
    
    with gr.Tab("Recomendaciones"):
        recom_btn = gr.Button("Generar Recomendaciones", variant="primary")
        recom_output = gr.Textbox(
            label="Recomendaciones para Mejora Continua", 
            placeholder="Las recomendaciones aparecerán aquí...",
            lines=15
        )
        
        with gr.Row():
            with gr.Column(scale=1):
                with gr.Group():
                    gr.Markdown("### Descargar Recomendaciones")
                    download_recom_txt = gr.File(label="Descargar como TXT")
                    download_recom_docx = gr.File(label="Descargar como DOCX")
    
    # Eventos
    caso_btn.click(
        fn=generar_caso_estudio,
        outputs=caso_output
    )
    
    plan_btn.click(
        fn=crear_planificacion,
        inputs=caso_output,
        outputs=plan_output
    )
    
    solucion_btn.click(
        fn=generar_solucion,
        inputs=caso_output,
        outputs=solucion_output
    ).then(
        lambda: gr.update(visible=True),
        None,
        [solucion_output]
    )
    
    comparar_btn.click(
        fn=comparar_respuestas,
        inputs=[respuesta_usuario, caso_output],
        outputs=[resultado_comparacion, grafico_comparacion]
    )
    
    recom_btn.click(
        fn=generar_recomendaciones,
        inputs=[respuesta_usuario, resultado_comparacion, caso_output],
        outputs=recom_output
    )
    
    # Eventos para descargar archivos
    caso_output.change(
        fn=lambda texto: generar_archivo_txt(texto, "caso_estudio"),
        inputs=caso_output,
        outputs=download_caso_txt
    ).then(
        fn=lambda texto: generar_archivo_docx(texto, "caso_estudio"),
        inputs=caso_output,
        outputs=download_caso_docx
    )
    
    plan_output.change(
        fn=lambda texto: generar_archivo_txt(texto, "planificacion"),
        inputs=plan_output,
        outputs=download_plan_txt
    ).then(
        fn=lambda texto: generar_archivo_docx(texto, "planificacion"),
        inputs=plan_output,
        outputs=download_plan_docx
    )
    
    solucion_output.change(
        fn=lambda texto: generar_archivo_txt(texto, "solucion"),
        inputs=solucion_output,
        outputs=download_sol_txt
    ).then(
        fn=lambda texto: generar_archivo_docx(texto, "solucion"),
        inputs=solucion_output,
        outputs=download_sol_docx
    )
    
    recom_output.change(
        fn=lambda texto: generar_archivo_txt(texto, "recomendaciones"),
        inputs=recom_output,
        outputs=download_recom_txt
    ).then(
        fn=lambda texto: generar_archivo_docx(texto, "recomendaciones"),
        inputs=recom_output,
        outputs=download_recom_docx
    )

# Iniciar la aplicación
if __name__ == "__main__":
    demo.launch(share=True)